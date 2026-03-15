"""
DOCX Parser
===========
Extracts text from Word documents using python-docx.
Preserves paragraph structure with newlines.
"""

from __future__ import annotations

from pathlib import Path

from backend.observability.logger import get_logger

logger = get_logger(__name__)


async def parse_docx(file_path: Path) -> str:
    """Extract all text from a DOCX file, preserving paragraph breaks."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    import asyncio

    def _extract() -> str:
        doc = Document(str(file_path))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        return "\n\n".join(paragraphs)

    loop = asyncio.get_event_loop()
    text = await loop.run_in_executor(None, _extract)

    logger.debug("docx_parsed", path=str(file_path), chars=len(text))
    return text